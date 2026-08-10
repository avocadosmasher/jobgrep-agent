"""이력서 파싱 — 텍스트 레이어 우선, OCR은 폴백 (T21, 설계도 §8-5).

이력서·경력기술서는 대부분 텍스트 레이어가 살아 있는 PDF/DOCX다. OCR을 무조건
태우면 품질이 오히려 나빠지므로 OCR은 기본 경로가 아니라 **폴백**이다.

    1) 확장자 분기      .docx → 문서 파서 / .pdf → 텍스트 레이어 / 이미지 → 3)
    2) 품질 검사        문자 수·공백 비율·판독 가능 문자 — **코드로, 결정론적**
    3) OCR 폴백         tools/ocr.py (T22). 없으면 조용히 건너뛴다
    4) 실패             건진 텍스트를 그대로 돌려주고 H4 보정에 맡긴다 (T22)

신뢰도는 **어느 층에서 나왔는지**로 정해진다 — `CONFIDENCE_BY_LAYER`가 유일한
정의 지점이다(T16 `tools/fetch_jd.py`와 같은 규약).

OCR 이음매 (T22가 읽을 것)
--------------------------
T22의 소유 파일은 `tools/ocr.py`·`app/hitl.py`뿐이라 **이 파일을 수정할 수 없다.**
그래서 호출부를 여기에 미리 뚫어 뒀다. T22는 아래 모양의 함수를 만들기만 하면
`parse_resume`이 자동으로 그 경로를 탄다 — 배선을 위해 이 파일로 돌아올 필요가 없다.

    # tools/ocr.py
    def ocr_document(file_path: str) -> str:
        \"\"\"이미지·스캔 PDF에서 텍스트를 추출한다. 실패는 예외가 아니라 빈 문자열.\"\"\"

모듈이 없으면 `ImportError`를 삼키고 4)로 떨어진다. 그래서 T21은 T22 없이 완결이며,
T22가 들어오는 순간 "만들어졌는데 아무도 안 부르는 모듈"(D39·D43·D59)이 되지 않는다.
"""

from __future__ import annotations

from pathlib import Path

from contracts.enums import Confidence

# --- 층 → 신뢰도. 유일한 정의 지점 (T16 CONFIDENCE_BY_LAYER와 같은 어휘) ---
LAYER_TEXT = "텍스트레이어"
LAYER_OCR = "OCR"
LAYER_FAILED = "실패"

CONFIDENCE_BY_LAYER: dict[str, Confidence] = {
    LAYER_TEXT: Confidence.HIGH,
    LAYER_OCR: Confidence.MID,
    LAYER_FAILED: Confidence.LOW,
}

# --- 확장자 분기 ---
PDF_SUFFIXES = frozenset({".pdf"})
DOCX_SUFFIXES = frozenset({".docx"})
IMAGE_SUFFIXES = frozenset(
    {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
)

# --- 품질 임계 (결정론적 — LLM 금지) ---
# 스캔본은 텍스트 레이어가 아예 없어 0자로 떨어진다. 임계를 이력서 한 장의
# 최소 분량보다 넉넉히 낮게 잡은 것은, 경계에서 애매하면 OCR을 태우는 쪽이
# 안전하기 때문이다 — 잘못 태우면 느려질 뿐이지만 잘못 통과시키면 빈 프로필이 된다.
MIN_CHARS = 200
MAX_WHITESPACE_RATIO = 0.5
MIN_READABLE_CHARS = 50

_HANGUL_RANGES = (
    ("가", "힣"),  # 음절
    ("ᄀ", "ᇿ"),  # 자모
    ("㄰", "㆏"),  # 호환 자모
)


def count_hangul(text: str) -> int:
    """한글 문자 수."""
    return sum(
        1 for c in text if any(lo <= c <= hi for lo, hi in _HANGUL_RANGES)
    )


def count_latin(text: str) -> int:
    """라틴 문자 수."""
    return sum(1 for c in text if c.isascii() and c.isalpha())


def whitespace_ratio(text: str) -> float:
    """공백 비율. 빈 문자열은 1.0 — 전부 쓸모없다는 뜻이다."""
    if not text:
        return 1.0
    return sum(1 for c in text if c.isspace()) / len(text)


def text_quality_ok(text: str) -> bool:
    """추출 텍스트가 쓸 만한지 판단한다 — **코드로, LLM 없이**.

    세 가지를 본다. ① 문자 수 ② 공백 비율 ③ 판독 가능 문자 수.

    ③을 "한글 포함 여부"가 아니라 "한글 **또는** 라틴 문자"로 둔 이유가 있다.
    이 검사가 가려내야 하는 것은 *언어*가 아니라 **깨진 인코딩**이다. 한글을
    필수로 걸면 영문 이력서가 통째로 스캔본 취급을 받아 멀쩡한 텍스트 레이어를
    버리고 OCR을 타게 된다. 반대로 CJK 폰트가 깨져 나온 모지바케는 한글도
    라틴도 아니라 이 조건에서 그대로 걸린다.
    """
    stripped = text.strip()
    if len(stripped) < MIN_CHARS:
        return False
    if whitespace_ratio(text) > MAX_WHITESPACE_RATIO:
        return False
    readable = max(count_hangul(text), count_latin(text))
    return readable >= MIN_READABLE_CHARS


def needs_manual_correction(text: str, confidence: Confidence) -> bool:
    """H4(사용자 직접 보정, T22)를 띄워야 하는지.

    OCR까지 태웠는데도 못 건진 경우다. 화면은 이 판별자만 보면 되고 층 이름을
    알 필요가 없다 — T16 `is_uncollected()`와 같은 역할이다.
    """
    return confidence is Confidence.LOW or not text.strip()


# --- 층별 추출기 ---
def _extract_pdf_text_layer(file_path: str) -> str:
    """PDF의 텍스트 레이어를 읽는다. 스캔본이면 빈 문자열에 가깝게 나온다."""
    import pymupdf

    try:
        with pymupdf.open(file_path) as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception:
        # 열리지 않는 PDF도 실패가 아니라 빈 결과다 — 상위에서 OCR·H4로 흐른다.
        return ""


def _extract_docx(file_path: str) -> str:
    """DOCX 본문을 읽는다. 표 안의 글자도 이력서에서는 본문이다."""
    import docx

    try:
        document = docx.Document(file_path)
    except Exception:
        return ""

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part.strip())


def _try_ocr(file_path: str) -> str:
    """T22의 OCR을 부른다. 아직 없으면 조용히 빈 문자열."""
    try:
        from tools.ocr import ocr_document  # T22가 만든다
    except ImportError:
        return ""

    try:
        return ocr_document(file_path) or ""
    except Exception:
        # OCR 실패를 전체 실패로 만들지 않는다 (T22 불변식) — H4로 흐른다.
        return ""


def parse_resume(file_path: str) -> tuple[str, Confidence]:
    """이력서 파일을 텍스트로 파싱하고 추출 신뢰도를 함께 반환한다.

    계약: `contracts/tools.py::parse_resume`.

    반환하는 신뢰도는 **어느 층에서 건졌는지**다 — 텍스트 레이어 `상`,
    OCR `중`, 실패 `하`. 실패해도 예외를 던지지 않고 **건진 텍스트를 그대로**
    돌려준다. 사용자가 화면에서 고칠 수 있어야 하기 때문이다(H4, T22).
    """
    suffix = Path(file_path).suffix.lower()

    # 1) 확장자 분기
    if suffix in DOCX_SUFFIXES:
        text = _extract_docx(file_path)
    elif suffix in PDF_SUFFIXES:
        text = _extract_pdf_text_layer(file_path)
    elif suffix in IMAGE_SUFFIXES:
        text = ""  # 이미지는 텍스트 레이어라는 게 없다 — 3)으로 직행
    else:
        # 모르는 확장자. 추측해서 파서를 태우지 않는다.
        return "", CONFIDENCE_BY_LAYER[LAYER_FAILED]

    # 2) 품질 검사 — 통과하면 OCR을 태우지 않는다
    if text_quality_ok(text):
        return text, CONFIDENCE_BY_LAYER[LAYER_TEXT]

    # 3) OCR 폴백 (T22)
    ocr_text = _try_ocr(file_path)
    if text_quality_ok(ocr_text):
        return ocr_text, CONFIDENCE_BY_LAYER[LAYER_OCR]

    # 4) 실패 — 둘 중 그나마 건진 쪽을 넘겨 H4 보정에 맡긴다
    best = ocr_text if len(ocr_text.strip()) > len(text.strip()) else text
    return best, CONFIDENCE_BY_LAYER[LAYER_FAILED]
